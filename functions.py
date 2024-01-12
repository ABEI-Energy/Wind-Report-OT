'''

This script contains the functions to be called from streamlit.py

'''
import pandas as pd
import locale as lc
import datetime as dt
import streamlit as st
import requests
import utm
from docx import Document
from PIL import Image
from docx.shared import Cm
import io


lc.setlocale(lc.LC_ALL,'en_US.UTF-8')

month = dt.datetime.now().strftime("%B %Y")
year = month.split(' ')[-1]

@st.cache_data
def textreaderFurow(file):

    dictFurow = dict()

    lines = file.getvalue().decode('utf_8').split('\r\n')

    lines = [line.strip().split('\t') for line in lines]

    # Get key values
    for i, row in enumerate(lines):
        if len(row)==2:
            dictFurow[row[0]] = row[1]

    # Get table 
    w1 = "General characteristics"
    for i, row in enumerate(lines):
        if w1 in row:
            i += 1
            break
    lines = lines[i:]
    df = pd.DataFrame(columns = lines[0])
    lines = lines[1:]
    for i, row in enumerate(lines):
        if len(row)!=1:
            df.loc[len(df)] = row
    
    df['Wake losses (%)'] = round(100 - df['Array Efficiency [%]'].apply(float),2)


    # Añadimos coordenadas de 

    dictFurow['wakeLoss'] = round(df['Array Efficiency [%]'].apply(float).mean(),2)

    # Calculamos el centroide del parque
    dictFurow['latUTMProj'] = round(df['Y [m]'].apply(float).sum()/df['Y [m]'].count(),2)
    dictFurow['longUTMProj'] = round(df['X [m]'].apply(float).sum()/df['X [m]'].count(),2) 

    # Get first dataframe
    df['Turbine ID'] = 'WT ' + df['Turbine ID']
    df1 = df.copy()
    df1 = df1[['Turbine ID','X [m]','Y [m]','Terrain Elevation [m]','Nearest Turbine ID','Distance to the nearest Turbine [m]']]

    df2 = df.copy()
    df2 = df2[['Turbine ID','Capacity [kW]','Mean Free [m/s]','Gross Yield [MWh]','Wake losses (%)','Net Yield [MWh]', 'Net Capacity factor [%]', 'Net Full load hours [h]']]

    return dictFurow, df, df1, df2

@st.cache_data
def dictMaker(dict, program, df):

    dict['dateTime'] = month
    dict['currentYear'] = year
    dict['yearMinus'] = str(int(float(year) - float(dict['numYears'])))
    dict['dateTable'] = dt.datetime.now().strftime("%d/%m/%y")
    
    
    if program == 'Furow':
        
        dict['turbulenceIntensity'] = dict.pop('Total TI [%]')
        dict['equivalentHours'] = dict.pop('Net Full load hours [h]')
        dict['husoUTM'] = dict.pop('UTM Zone')
        dict['powerWF'] = dict.pop('Total Capacity Installed [MW]')
        dict['capacityFactor'] = dict.pop('Net Capacity Factor [%]')
        dict['anualProd'] = dict.pop('Net Yield [MWh]')
        dict['anualProdGross'] = dict.pop('Ideal Yield [MWh]')
        dict['numTurbines'] = dict.pop('Number of turbines')
        dict['wakeSpeed'] = dict.pop('Mean Wake Affected Wind Speed [m/s]')
        dict['windSpeed'] = dict.pop('Mean Free Wind Speed [m/s]')
        dict['turbineModel'] = df['Turbine Type'].unique()[0]
        dict['hubHeight'] = df['Hub Height [m]'].unique()[0]
        dict['rotorSize'] = df['Rotor Diameter [m]'].unique()[0]
        dict['turbinePower'] = df['Capacity [kW]'].unique()[0]
        dict.pop('Wind Farm')
        dict['powerWF'] = dict['powerWF'].strip('0')
        dict['turbinePower'] = str(float(dict['turbinePower'].strip('0'))/1000)


        utm_to_latLon(dict)

        get_elevation(dict)

        dict.update(dict)

        st.session_state.dict = True

        return True


    else:
        pass
    
@st.cache_data
def initials(user):
    user = user.strip().split(' ')
    auxUser = []
    for word in user:
        auxUser.append([*word][0])
    user = '.'.join(auxUser)
    return user

@st.cache_data
def normalize(string):
    return str(string.replace(",", "."))

@st.cache_data
def get_elevation(dict):

    lat = dict['latProj']
    lon = dict['longProj']
    url = ('https://api.opentopodata.org/v1/test-dataset'f'?locations={lat},{lon}')
    while True:
        try:
           response = requests.get(url).json()
        except Exception as e:
           continue
        break
   
    elevation = response['results'][0]['elevation']

    dict['altProj'] = round(elevation, 2)

@st.cache_data
def utm_to_latLon(dict):
    lat_utm = float(dict['latUTMProj'])
    long_utm = float(dict['longUTMProj'])
    huso = int(dict['husoUTM'].split(' ')[0])
    zone = dict['husoUTM'].split(' ')[-1].strip()
    coordinates = utm.to_latlon(long_utm, lat_utm, huso, zone)
    dict['latProj'] = round(coordinates[0],2)
    dict['longProj'] = round(coordinates[-1],2)
    dict.update(dict)

@st.cache_data
def countryValues():
    df_countries = pd.read_csv('countries.csv', index_col=None).sort_values(by='Country')
    return df_countries

@st.cache_data
def ms_reader(directory):
    
    if 'csv' in directory:
        df = pd.read_csv(directory, index_col=None)

    elif 'xlsx' in directory:
        pass

    return df

@st.cache_data
def model_reader(directory):
    
    xls = pd.ExcelFile(directory)

    df_data = pd.read_excel(xls, 'Sheet1', index_col=None, header = None)


    df_power = pd.read_excel(xls, 'Sheet3', index_col=None, header = None)
    df_thrust = pd.read_excel(xls, 'Sheet4', index_col=None, header = None)

    df_power.drop([0,1,2,3], inplace = True)
    df_power.drop([0], axis = 1, inplace = True)
    df_power.rename(columns = df_power.iloc[0].apply(str), inplace = True)
    df_power.reset_index(inplace = True, drop = True)
    df_power.drop([0], inplace = True)
    df_power.set_index(df_power.iloc[:, 0].apply(str), inplace = True)
    df_power.index.names = ['Wind Speed']
    df_power.drop(df_power.columns[0], axis = 1, inplace = True)
    df_power.reset_index(inplace = True)
    df_power.sort_values(by = 'Wind Speed')
    
    df_power = df_power.astype(float)

    df_thrust.drop([0,1,2,3], inplace = True)
    df_thrust.drop([0], axis = 1, inplace = True)
    df_thrust.rename(columns = df_thrust.iloc[0].apply(str), inplace = True)
    df_thrust.reset_index(inplace = True, drop = True)
    df_thrust.drop([0], inplace = True)
    df_thrust.set_index(df_thrust.iloc[:, 0].apply(str), inplace = True)
    df_thrust.index.names = ['Wind Speed']
    df_thrust.drop(df_thrust.columns[0], axis = 1, inplace = True)

    df_thrust.reset_index(inplace = True)
    df_thrust.sort_values(by = 'Wind Speed')

    df_thrust = df_thrust.astype(float)


    # We have to choose the columns with 1.225, and would they not exist, interpolate them

    if '1.225' not in df_power.columns.values:
        ct_values = df_power.columns.values[1:].astype(float)
        i = 0
        for element in ct_values:
            if element < 1.225:
                i += 1
                continue
            else:
                minor_ct = ct_values[i]
                major_ct = ct_values[i+1]
        
        df_power['1.225'] = df_power[[str(minor_ct), str(major_ct)]].mean(axis=1)

    if '1.225' not in df_thrust.columns.values:
        ct_values = df_thrust.columns.values[1:].astype(float)
        i = 0
        for element in ct_values:
            if element < 1.225:
                i += 1
                continue
            else:
                minor_ct = ct_values[i]
                major_ct = ct_values[i+1]
        
        df_thrust['1.225'] = df_thrust[[str(minor_ct), str(major_ct)]].mean(axis=1)

    return df_power, df_thrust

def duplicateDoc():

    filemodelo = "resources/model/Wind report template.docx"

    return Document(filemodelo)

def resize_image(image, max_width):
    width, height = image.size
    if width > max_width:
        ratio = max_width / width
        new_width = max_width
        new_height = int(height * ratio)
        return image.resize((new_width, new_height))
    return image

def insert_image_in_cell(doc, picDict):

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if ("aerialCut" in cell.text):
                    imagen = Image.open(picDict['location'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 16.55:
                        imHeight = imHeight*16.55/imWidth
                        imWidth = 16.55
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))
                if ("powerCurvePic" in cell.text):
                    imagen = Image.open(picDict['powerCurvePic'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 16.55:
                        imHeight = imHeight*16.55/imWidth
                        imWidth = 16.55
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))
                if ("layWFout" in cell.text):
                    imagen = Image.open(picDict['layout'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 16.55:
                        imHeight = imHeight*16.55/imWidth
                        imWidth = 16.55
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))  
                if ("windHeatMap" in cell.text):
                    imagen = Image.open(picDict['wind resource'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 16.55:
                        imHeight = imHeight*16.55/imWidth
                        imWidth = 16.55
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))  
                if ("turbulenceHeatMap" in cell.text):
                    imagen = Image.open(picDict['turbulence'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 16.55:
                        imHeight = imHeight*16.55/imWidth
                        imWidth = 16.55
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))                                                           
    
    st.session_state.picsDone = True
    
def docWriter(docxFile,docxDict):

    #Headers
    for section in docxFile.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word in docxDict:
                            if word in paragraph.text:
                                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                                paragraph.style = docxFile.styles['headerBlue'] 


    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # previousStyle = paragraph.style
                    for word in docxDict:
                        if word in paragraph.text:
                            previousStyle = paragraph.style.name
                            paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                            paragraph.style = docxFile.styles[previousStyle]   

    
        #Resto del documento
    for paragraph in docxFile.paragraphs:
        for word in docxDict:
            if word in paragraph.text:
                previousStyle = paragraph.style.name
                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                paragraph.style = docxFile.styles[previousStyle] 


    st.session_state.wordsDone = True
    st.session_state.tablesDone = True

def docTabler(docxFile, df1, df2):
    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                if "flagDf1" in cell.text:
                    cell.text = "turbine ID"
                    for i in range(len(df1)):
                        table.add_row()
                        table.style = 'tableBlue'

                    for i in range(df1.shape[0]):
                        for j in range(df1.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df1.values[i,j])

                    table.style = 'tableBlue'

                if "flagDf2" in cell.text:
                    cellStyle = cell.paragraphs[0].style.name
                    cell.text = "turbine ID"
                    for i in range(len(df2)):
                        table.add_row()
                        table.style = 'tableBlue'

                    for i in range(df2.shape[0]):
                        for j in range(df2.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df2.values[i,j])

                    table.style = 'tableBlue'

    st.session_state.tablerDone = True

















